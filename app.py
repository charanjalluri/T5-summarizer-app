"""
T5 Supercharged Summarization Tool
Multi-modal AI summarization with authentication
"""

import os
import io
import tempfile
from typing import List, Dict, Optional
import hashlib

import streamlit as st
from PIL import Image
import numpy as np

# PDF handling
try:
    import fitz  # PyMuPDF
    PYMUPDF_OK = True
except:
    PYMUPDF_OK = False

try:
    import pdfplumber
    PDFPLUMBER_OK = True
except:
    PDFPLUMBER_OK = False

from PyPDF2 import PdfReader

# Audio
import speech_recognition as sr
from pydub import AudioSegment

# NLP
import langdetect
from deep_translator import GoogleTranslator

# Transformers
try:
    import torch
    from transformers import T5Tokenizer, T5ForConditionalGeneration
except:
    torch = None
    T5Tokenizer = None

# Embeddings
try:
    from sentence_transformers import SentenceTransformer
except:
    SentenceTransformer = None

from sklearn.metrics.pairwise import cosine_similarity

try:
    import faiss
    FAISS_OK = True
except:
    FAISS_OK = False

# Graph
import networkx as nx
from pyvis.network import Network
from streamlit.components.v1 import html as st_html

# Exports
from reportlab.pdfgen import canvas as rl_canvas
from docx import Document
from gtts import gTTS

# Gemini
from google import genai
from google.genai import types

# ================================
# CONFIGURATION
# ================================

st.set_page_config(
    page_title="🚀 AI Summarizer Pro",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ================================
# AUTHENTICATION SYSTEM
# ================================

# Simple user database (in production, use a real database)
USERS = {
    "admin": hashlib.sha256("admin123".encode()).hexdigest(),
    "demo": hashlib.sha256("demo123".encode()).hexdigest(),
}

def check_password():
    """Returns True if user is authenticated"""
    
    def login_form():
        with st.form("login_form"):
            st.subheader("🔐 Login Required")
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submit = st.form_submit_button("Login")
            
            if submit:
                password_hash = hashlib.sha256(password.encode()).hexdigest()
                if username in USERS and USERS[username] == password_hash:
                    st.session_state["authenticated"] = True
                    st.session_state["username"] = username
                    st.success("✅ Login successful!")
                    st.rerun()
                else:
                    st.error("❌ Invalid credentials")
                    st.info("**Demo credentials:**\nUsername: `demo`\nPassword: `demo123`")
    
    # Check if already authenticated
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False
    
    if not st.session_state["authenticated"]:
        login_form()
        return False
    
    return True

def logout():
    """Logout current user"""
    st.session_state["authenticated"] = False
    st.session_state["username"] = None
    st.rerun()

# ================================
# GEMINI CLIENT
# ================================

def get_gemini_client():
    """Initialize Gemini client with API key"""
    # Priority: Streamlit secrets > Environment > User input
    api_key = None
    
    if hasattr(st, 'secrets') and 'GEMINI_API_KEY' in st.secrets:
        api_key = st.secrets['GEMINI_API_KEY']
    elif 'GEMINI_API_KEY' in os.environ:
        api_key = os.environ['GEMINI_API_KEY']
    
    if not api_key:
        with st.sidebar:
            api_key = st.text_input("🔑 Gemini API Key", type="password", 
                                   help="Get free key at https://aistudio.google.com/apikey")
    
    if api_key:
        try:
            return genai.Client(api_key=api_key)
        except Exception as e:
            st.error(f"Failed to initialize Gemini: {e}")
            return None
    return None

# ================================
# CACHED RESOURCES
# ================================

@st.cache_resource(show_spinner=False)
def load_t5(model_name: str):
    """Load T5 model (fallback option)"""
    if T5Tokenizer is None:
        raise RuntimeError("Transformers not available")
    tokenizer = T5Tokenizer.from_pretrained(model_name)
    model = T5ForConditionalGeneration.from_pretrained(model_name)
    device = "cuda" if (torch and torch.cuda.is_available()) else "cpu"
    if device == "cuda":
        model.to("cuda")
    return tokenizer, model, device

@st.cache_resource(show_spinner=False)
def load_embedder():
    """Load sentence embedder"""
    if SentenceTransformer is None:
        raise RuntimeError("sentence-transformers not available")
    return SentenceTransformer("all-MiniLM-L6-v2")

@st.cache_resource(show_spinner=False)
def load_spacy():
    """Load spaCy for NER"""
    try:
        import spacy
        return spacy.load("en_core_web_sm")
    except:
        return None

# ================================
# UTILITY FUNCTIONS
# ================================

def detect_language(text: str) -> str:
    """Detect text language"""
    try:
        return langdetect.detect(text)
    except:
        return "en"

def translate_text(text: str, target_lang: str = "en") -> str:
    """Translate text to target language"""
    if not text:
        return text
    try:
        return GoogleTranslator(source="auto", target=target_lang).translate(text)
    except:
        return text

def chunk_text(text: str, chunk_size: int = 1200, overlap: int = 120) -> List[str]:
    """Split text into overlapping chunks"""
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = words[i:i + chunk_size]
        chunks.append(" ".join(chunk))
        i += max(1, chunk_size - overlap)
    return chunks

# ================================
# PDF EXTRACTION
# ================================

def extract_text_from_pdf(file) -> Dict:
    """Extract text from PDF using multiple methods"""
    text_all = []
    tables = []
    pages = 0
    
    file_bytes = file.read()
    
    # Try PyMuPDF first (best)
    if PYMUPDF_OK:
        try:
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                pages = doc.page_count
                for page in doc:
                    text_all.append(page.get_text("text") or "")
        except Exception as e:
            st.info(f"PyMuPDF fallback: {e}")
    
    # Try pdfplumber
    if PDFPLUMBER_OK and not text_all:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = len(pdf.pages)
                for p in pdf.pages:
                    t = p.extract_text() or ""
                    if t:
                        text_all.append(t)
                    try:
                        tbs = p.extract_tables() or []
                        for tb in tbs:
                            if tb:
                                tables.append(tb)
                    except:
                        pass
        except Exception as e:
            st.info(f"pdfplumber fallback: {e}")
    
    # Fallback to PyPDF2
    if not text_all:
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            pages = len(reader.pages)
            for page in reader.pages:
                text_all.append(page.extract_text() or "")
        except Exception as e:
            return {"text": "", "pages": 0, "tables": [], "error": str(e)}
    
    return {
        "text": "\n".join(text_all),
        "pages": pages,
        "tables": tables
    }

# ================================
# AUDIO TRANSCRIPTION
# ================================

def transcribe_audio(uploaded_audio) -> str:
    """Convert audio to text"""
    try:
        audio = AudioSegment.from_file(io.BytesIO(uploaded_audio.read()))
        buf = io.BytesIO()
        audio.export(buf, format="wav")
        
        recognizer = sr.Recognizer()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(buf.getvalue())
            tmp.flush()
            
            with sr.AudioFile(tmp.name) as source:
                audio_data = recognizer.record(source)
                try:
                    return recognizer.recognize_google(audio_data)
                except:
                    return ""
    except Exception as e:
        st.error(f"Transcription failed: {e}")
        return ""

# ================================
# SUMMARIZATION FUNCTIONS
# ================================

def gemini_summarize(client, text: str, style: str = "concise", target_lang: str = "en") -> str:
    """Summarize using Gemini"""
    if not client or not text.strip():
        return ""
    
    style_prompts = {
        "concise": "Create a brief summary in 3-5 sentences.",
        "bullet": "Create a bullet-point summary with key takeaways.",
        "detailed": "Create a comprehensive summary with main points and details.",
    }
    
    prompt = f"{style_prompts.get(style, style_prompts['concise'])}\n\nLanguage: {target_lang}\n\n{text[:8000]}"
    
    try:
        response = client.models.generate_content(
            model="gemini-1.5-flash",
            contents=prompt
        )
        return response.text or ""
    except Exception as e:
        st.error(f"Gemini error: {e}")
        return ""

def gemini_qa(client, question: str, context: str, target_lang: str = "en") -> str:
    """Answer questions using Gemini"""
    if not client:
        return ""
    
    prompt = f"Context:\n{context[:6000]}\n\nQuestion: {question}\n\nAnswer in {target_lang}:"
    
    try:
        response = client.models.generate_content(
            model="gemini-1.5-flash",
            contents=prompt
        )
        return response.text or ""
    except Exception as e:
        st.error(f"Q&A error: {e}")
        return ""

def gemini_vision(client, image_bytes: bytes, mime_type: str, target_lang: str = "en") -> str:
    """Analyze image using Gemini Vision"""
    if not client:
        return ""
    
    try:
        response = client.models.generate_content(
            model="gemini-1.5-flash",
            contents=[
                types.Part(inline_data=types.Blob(mime_type=mime_type, data=image_bytes)),
                types.Part(text=f"Describe and extract text. Language: {target_lang}")
            ]
        )
        return response.text or ""
    except Exception as e:
        st.error(f"Vision error: {e}")
        return ""

# ================================
# EXPORT FUNCTIONS
# ================================

def save_as_txt(text: str) -> bytes:
    """Export as TXT"""
    return text.encode("utf-8")

def save_as_pdf(text: str) -> bytes:
    """Export as PDF"""
    buffer = io.BytesIO()
    c = rl_canvas.Canvas(buffer)
    c.setFont("Helvetica", 12)
    
    y = 800
    for line in text.split("\n"):
        if y < 50:
            c.showPage()
            y = 800
        c.drawString(50, y, line[:90])
        y -= 15
    
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def save_as_docx(text: str) -> bytes:
    """Export as DOCX"""
    doc = Document()
    doc.add_heading("Summary", level=1)
    for para in text.split("\n\n"):
        doc.add_paragraph(para)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ================================
# MAIN APP
# ================================

def main():
    # Check authentication
    if not check_password():
        return
    
    # Header with logout
    col1, col2 = st.columns([6, 1])
    with col1:
        st.title("🚀 AI Summarizer Pro")
        st.caption(f"Welcome, **{st.session_state.get('username', 'User')}**!")
    with col2:
        if st.button("🚪 Logout", use_container_width=True):
            logout()
    
    st.markdown("---")
    
    # Initialize Gemini
    client = get_gemini_client()
    if not client:
        st.warning("⚠️ Please provide Gemini API Key to use AI features")
        st.info("Get free API key: https://aistudio.google.com/apikey")
        return
    
    # Sidebar settings
    with st.sidebar:
        st.header("⚙️ Settings")
        input_type = st.selectbox("📥 Input Type", 
                                  ["Text", "PDF", "Image", "Audio"])
        style = st.selectbox("✍️ Summary Style", 
                            ["concise", "bullet", "detailed"])
        target_lang = st.text_input("🌐 Language Code", "en",
                                   help="e.g., en, es, fr, hi")
        
        st.markdown("---")
        st.markdown("### 📊 Usage Stats")
        if "summary_count" not in st.session_state:
            st.session_state.summary_count = 0
        st.metric("Summaries Generated", st.session_state.summary_count)
    
    # Main content area
    if input_type == "Text":
        handle_text_input(client, style, target_lang)
    elif input_type == "PDF":
        handle_pdf_input(client, style, target_lang)
    elif input_type == "Image":
        handle_image_input(client, style, target_lang)
    elif input_type == "Audio":
        handle_audio_input(client, style, target_lang)

def handle_text_input(client, style, target_lang):
    """Handle text summarization"""
    st.subheader("📝 Text Summarization")
    
    text = st.text_area("Enter or paste text:", height=200,
                       placeholder="Paste your text here...")
    
    if st.button("✨ Summarize", type="primary"):
        if not text.strip():
            st.warning("Please enter some text")
            return
        
        with st.spinner("Generating summary..."):
            summary = gemini_summarize(client, text, style, target_lang)
            
            if summary:
                st.session_state.summary_count += 1
                st.success("✅ Summary generated!")
                
                st.markdown("### 📄 Summary")
                st.write(summary)
                
                # Download buttons
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.download_button("📥 TXT", save_as_txt(summary), 
                                     "summary.txt", "text/plain")
                with col2:
                    st.download_button("📥 PDF", save_as_pdf(summary),
                                     "summary.pdf", "application/pdf")
                with col3:
                    st.download_button("📥 DOCX", save_as_docx(summary),
                                     "summary.docx", 
                                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

def handle_pdf_input(client, style, target_lang):
    """Handle PDF summarization"""
    st.subheader("📄 PDF Summarization")
    
    uploaded_pdf = st.file_uploader("Upload PDF", type="pdf")
    
    if uploaded_pdf:
        with st.spinner("Extracting PDF..."):
            data = extract_text_from_pdf(uploaded_pdf)
        
        if data.get("error"):
            st.error(f"Error: {data['error']}")
            return
        
        pdf_text = data.get("text", "")
        pages = data.get("pages", 0)
        
        st.success(f"✅ Extracted {pages} pages")
        
        with st.expander("📖 Preview"):
            st.text(pdf_text[:1000] + "...")
        
        if st.button("✨ Summarize PDF", type="primary"):
            with st.spinner("Summarizing..."):
                summary = gemini_summarize(client, pdf_text, style, target_lang)
                
                if summary:
                    st.session_state.summary_count += 1
                    st.markdown("### 📄 Summary")
                    st.write(summary)
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button("📥 TXT", save_as_txt(summary),
                                         "summary.txt")
                    with col2:
                        st.download_button("📥 PDF", save_as_pdf(summary),
                                         "summary.pdf")

def handle_image_input(client, style, target_lang):
    """Handle image analysis"""
    st.subheader("🖼️ Image Analysis")
    
    uploaded_image = st.file_uploader("Upload Image", 
                                     type=["png", "jpg", "jpeg"])
    
    if uploaded_image:
        image = Image.open(uploaded_image)
        st.image(image, caption="Uploaded Image", use_container_width=True)
        
        if st.button("🔍 Analyze", type="primary"):
            with st.spinner("Analyzing image..."):
                result = gemini_vision(client, uploaded_image.getvalue(),
                                     uploaded_image.type, target_lang)
                
                if result:
                    st.session_state.summary_count += 1
                    st.markdown("### 📊 Analysis")
                    st.write(result)
                    
                    st.download_button("📥 Save", save_as_txt(result),
                                     "analysis.txt")

def handle_audio_input(client, style, target_lang):
    """Handle audio transcription"""
    st.subheader("🎤 Audio Transcription")
    
    uploaded_audio = st.file_uploader("Upload Audio",
                                     type=["wav", "mp3", "m4a"])
    
    if uploaded_audio:
        st.audio(uploaded_audio)
        
        if st.button("🎯 Transcribe & Summarize", type="primary"):
            with st.spinner("Transcribing..."):
                text = transcribe_audio(uploaded_audio)
                
                if text:
                    st.markdown("### 📝 Transcript")
                    st.write(text)
                    
                    with st.spinner("Summarizing..."):
                        summary = gemini_summarize(client, text, style, target_lang)
                        
                        if summary:
                            st.session_state.summary_count += 1
                            st.markdown("### 📄 Summary")
                            st.write(summary)
                            
                            st.download_button("📥 Save", save_as_txt(summary),
                                             "summary.txt")
                else:
                    st.warning("No speech detected")

if __name__ == "__main__":
    main()
